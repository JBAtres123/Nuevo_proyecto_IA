"""
EduGrade AI — Servicio de Generación de Reportes
Genera informes profesionales en PDF (fpdf2) y Word (python-docx).
Mantiene soporte para LaTeX en texto usando notación textual.
"""
from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path

from config import get_settings

settings = get_settings()


def _latex_to_text(text: str) -> str:
    """
    Convierte expresiones LaTeX básicas a texto legible para PDF/Word.
    Para una renderización real se necesitaría una librería como matplotlib.
    """
    if not text:
        return text
    # Fracciones
    text = re.sub(r"\\frac\{(.+?)\}\{(.+?)\}", r"(\1)/(\2)", text)
    # Raíces
    text = re.sub(r"\\sqrt\[(.+?)\]\{(.+?)\}", r"\1√(\2)", text)
    text = re.sub(r"\\sqrt\{(.+?)\}", r"√(\1)", text)
    # Superíndices/subíndices
    text = re.sub(r"\^\{(.+?)\}", r"^\1", text)
    text = re.sub(r"_\{(.+?)\}", r"_\1", text)
    # Remover delimitadores $
    text = re.sub(r"\$\$(.+?)\$\$", r"[\1]", text, flags=re.DOTALL)
    text = re.sub(r"\$(.+?)\$", r"\1", text)
    # Comandos comunes
    text = text.replace(r"\cdot", "·").replace(r"\times", "×")
    text = text.replace(r"\div", "÷").replace(r"\pm", "±")
    text = text.replace(r"\leq", "≤").replace(r"\geq", "≥")
    text = text.replace(r"\neq", "≠").replace(r"\approx", "≈")
    text = text.replace(r"\infty", "∞").replace(r"\alpha", "α")
    text = text.replace(r"\beta", "β").replace(r"\pi", "π")
    return text


def generate_pdf_report(resultado: dict, output_dir: str | None = None) -> str:
    """
    Genera un informe PDF profesional de la calificación.
    
    Args:
        resultado: dict con la calificación final del agente Juez
        output_dir: directorio de salida (default: settings.reports_dir)
    
    Returns:
        Ruta del archivo PDF generado
    """
    from fpdf import FPDF

    output_dir = output_dir or settings.reports_dir
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    estudiante = resultado.get("estudiante", "Estudiante")
    serie = resultado.get("serie", "")
    punteo = resultado.get("punteo_total", 0)
    maximo = resultado.get("punteo_maximo_total", 100)
    porcentaje = resultado.get("porcentaje", 0)
    nivel = resultado.get("nivel_aplicado", 5)

    class PDF(FPDF):
        def header(self):
            self.set_fill_color(30, 58, 138)   # azul oscuro
            self.rect(0, 0, 210, 22, "F")
            self.set_text_color(255, 255, 255)
            self.set_font("Helvetica", "B", 14)
            self.cell(0, 12, "  EduGrade AI — Informe de Calificación", ln=False, align="L")
            self.set_font("Helvetica", "", 9)
            self.cell(0, 12, datetime.now().strftime("%d/%m/%Y %H:%M  "), ln=True, align="R")
            self.set_text_color(0, 0, 0)
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"Generado por EduGrade AI  |  Página {self.page_no()}", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # ── Datos del estudiante ─────────────────────────────────
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(30, 58, 138)
    pdf.cell(0, 8, "Datos del Estudiante", ln=True)
    pdf.set_draw_color(30, 58, 138)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(3)

    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 11)
    col_w = 45
    data_rows = [
        ("Estudiante:", estudiante),
        ("Serie:", serie or "N/A"),
        ("Nivel de exigencia:", f"{nivel}/10"),
    ]
    for label, value in data_rows:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(col_w, 7, label)
        pdf.set_font("Helvetica", "", 11)
        pdf.cell(0, 7, str(value), ln=True)

    # ── Punteo destacado ──────────────────────────────────────
    pdf.ln(3)
    color = (22, 163, 74) if porcentaje >= 70 else (234, 179, 8) if porcentaje >= 60 else (220, 38, 38)
    pdf.set_fill_color(*color)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 12,
             f"  PUNTEO FINAL: {punteo}/{maximo} pts  —  {porcentaje:.1f}%",
             ln=True, fill=True, align="L")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(5)

    # ── Detalle por pregunta ──────────────────────────────────
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(30, 58, 138)
    pdf.cell(0, 8, "Detalle por Pregunta", ln=True)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(3)

    for p in resultado.get("preguntas", []):
        num = p.get("numero", "?")
        serie_sec = p.get("serie_seccion", "")
        pob = p.get("punteo_obtenido", 0)
        pmax = p.get("punteo_maximo", 0)
        correcta = p.get("es_correcta", False)
        just = _latex_to_text(p.get("justificacion", ""))
        resp_est = _latex_to_text(p.get("respuesta_estudiante", ""))

        # Encabezado de pregunta
        icon = "✓" if correcta else "✗"
        pdf.set_fill_color(240, 249, 255) if correcta else pdf.set_fill_color(255, 245, 245)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(30, 58, 138)
        pdf.cell(0, 7,
                 f"  {icon}  Pregunta {num} {('— Sección ' + serie_sec) if serie_sec else ''}   "
                 f"[{pob}/{pmax} pts]",
                 ln=True, fill=True)
        pdf.set_text_color(0, 0, 0)

        if resp_est:
            pdf.set_font("Helvetica", "I", 10)
            pdf.set_x(15)
            pdf.multi_cell(180, 5, f"Respuesta: {resp_est[:200]}")

        if just:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_x(15)
            pdf.multi_cell(180, 5, f"Observación: {just}")
        pdf.ln(2)

    # ── Conclusión ────────────────────────────────────────────
    pdf.ln(2)
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(30, 58, 138)
    pdf.cell(0, 8, "Conclusión General", ln=True)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(3)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, resultado.get("conclusion", "Sin conclusión."))

    # ── Fortalezas / Debilidades / Sugerencias ────────────────
    sections = [
        ("Fortalezas", resultado.get("fortalezas", []), (22, 163, 74)),
        ("Áreas de Mejora", resultado.get("debilidades", []), (220, 38, 38)),
        ("Sugerencias", resultado.get("sugerencias", []), (234, 179, 8)),
    ]
    for title, items, color in sections:
        if not items:
            continue
        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(*color)
        pdf.cell(0, 7, title, ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 11)
        for item in items:
            pdf.set_x(15)
            pdf.multi_cell(180, 5.5, f"• {_latex_to_text(str(item))}")

    # ── Guardar ───────────────────────────────────────────────
    safe_name = re.sub(r"[^\w\-]", "_", estudiante)[:40]
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"informe_{safe_name}_{ts}.pdf"
    output_path = str(Path(output_dir) / filename)
    pdf.output(output_path)
    return output_path


def generate_word_report(resultado: dict, output_dir: str | None = None) -> str:
    """
    Genera un informe Word (.docx) profesional de la calificación.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    output_dir = output_dir or settings.reports_dir
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    estudiante = resultado.get("estudiante", "Estudiante")
    serie = resultado.get("serie", "")
    punteo = resultado.get("punteo_total", 0)
    maximo = resultado.get("punteo_maximo_total", 100)
    porcentaje = resultado.get("porcentaje", 0)
    nivel = resultado.get("nivel_aplicado", 5)

    doc = Document()

    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # ── Título ────────────────────────────────────────────────
    title = doc.add_heading("EduGrade AI — Informe de Calificación", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    sub = doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # ── Datos del estudiante ─────────────────────────────────
    doc.add_heading("Datos del Estudiante", 1).runs[0].font.color.rgb = RGBColor(30, 58, 138)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    cells_data = [
        ("Estudiante", estudiante),
        ("Serie", serie or "N/A"),
        ("Nivel de exigencia", f"{nivel}/10"),
    ]
    for i, (k, v) in enumerate(cells_data):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(v)

    doc.add_paragraph()

    # Punteo final
    p_punteo = doc.add_paragraph()
    run = p_punteo.add_run(f"PUNTEO FINAL: {punteo}/{maximo} pts  —  {porcentaje:.1f}%")
    run.bold = True
    run.font.size = Pt(14)
    if porcentaje >= 70:
        run.font.color.rgb = RGBColor(22, 163, 74)
    elif porcentaje >= 60:
        run.font.color.rgb = RGBColor(234, 179, 8)
    else:
        run.font.color.rgb = RGBColor(220, 38, 38)

    doc.add_paragraph()

    # ── Detalle por pregunta ──────────────────────────────────
    doc.add_heading("Detalle por Pregunta", 1).runs[0].font.color.rgb = RGBColor(30, 58, 138)

    for p in resultado.get("preguntas", []):
        num = p.get("numero", "?")
        serie_sec = p.get("serie_seccion", "")
        pob = p.get("punteo_obtenido", 0)
        pmax = p.get("punteo_maximo", 0)
        correcta = p.get("es_correcta", False)
        just = _latex_to_text(p.get("justificacion", ""))
        resp_est = _latex_to_text(p.get("respuesta_estudiante", ""))

        icon = "✓" if correcta else "✗"
        h = doc.add_heading(
            f"{icon} Pregunta {num}{(' — Sección ' + serie_sec) if serie_sec else ''}  [{pob}/{pmax} pts]",
            level=2
        )
        run = h.runs[0]
        run.font.color.rgb = RGBColor(22, 163, 74) if correcta else RGBColor(220, 38, 38)

        if resp_est:
            rp = doc.add_paragraph()
            rp.add_run("Respuesta: ").italic = True
            rp.add_run(resp_est[:300])

        if just:
            jp = doc.add_paragraph()
            jp.add_run("Observación: ").bold = True
            jp.add_run(just)

    # ── Conclusión ────────────────────────────────────────────
    doc.add_heading("Conclusión General", 1).runs[0].font.color.rgb = RGBColor(30, 58, 138)
    doc.add_paragraph(resultado.get("conclusion", "Sin conclusión."))

    # ── Fortalezas / Debilidades / Sugerencias ────────────────
    sections = [
        ("Fortalezas", resultado.get("fortalezas", []), RGBColor(22, 163, 74)),
        ("Áreas de Mejora", resultado.get("debilidades", []), RGBColor(220, 38, 38)),
        ("Sugerencias de Mejora", resultado.get("sugerencias", []), RGBColor(234, 179, 8)),
    ]
    for title, items, color in sections:
        if not items:
            continue
        h = doc.add_heading(title, 2)
        h.runs[0].font.color.rgb = color
        for item in items:
            doc.add_paragraph(f"• {_latex_to_text(str(item))}", style="List Bullet")

    # ── Guardar ───────────────────────────────────────────────
    safe_name = re.sub(r"[^\w\-]", "_", estudiante)[:40]
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"informe_{safe_name}_{ts}.docx"
    output_path = str(Path(output_dir) / filename)
    doc.save(output_path)
    return output_path


# Generador Word presentable. La ultima definicion de generate_word_report
# delega a este helper para evitar que versiones antiguas lo sobrescriban.
def _generate_word_report_pretty(resultado: dict, output_dir: str | None = None) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

    output_dir = output_dir or settings.reports_dir
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    estudiante = _clean_text(resultado.get("estudiante") or "Estudiante")
    serie_general = _clean_text(resultado.get("serie") or "Sin serie")
    examen_id = resultado.get("examen_id")
    fecha = resultado.get("fecha") or datetime.now().strftime("%d/%m/%Y, %I:%M %p")
    nivel = resultado.get("nivel_aplicado", 5)
    punteo = _num(resultado.get("punteo_total"))
    maximo = _num(resultado.get("punteo_maximo_total"))
    porcentaje = _num(resultado.get("porcentaje"))
    series = _series_agrupadas(resultado.get("preguntas", []))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9)

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = header.rows[0].cells
    for cell in (left, right):
        _shade_cell(cell, "0F172A")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    title = left.paragraphs[0]
    run = title.add_run(estudiante)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(248, 250, 252)

    meta = left.add_paragraph()
    meta_text = serie_general
    if examen_id:
        meta_text += f" - Examen #{examen_id}"
    meta_text += f" - {fecha}"
    meta_run = meta.add_run(meta_text)
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = RGBColor(203, 213, 225)

    score = right.paragraphs[0]
    score.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    score_run = score.add_run(f"{porcentaje:.1f}%")
    score_run.bold = True
    score_run.font.size = Pt(22)
    score_run.font.color.rgb = RGBColor(74, 222, 128) if porcentaje >= 70 else RGBColor(248, 113, 113)

    score_pts = right.add_paragraph()
    score_pts.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pts_run = score_pts.add_run(f"{_fmt_score(punteo)} / {_fmt_score(maximo)} pts")
    pts_run.font.size = Pt(8)
    pts_run.font.color.rgb = RGBColor(226, 232, 240)

    badges = doc.add_paragraph()
    badges.paragraph_format.space_before = Pt(8)
    b1 = badges.add_run("  completado  ")
    b1.bold = True
    b1.font.color.rgb = RGBColor(22, 101, 52)
    badges.add_run("   ")
    b2 = badges.add_run(f"  Exigencia {nivel}/10  ")
    b2.bold = True
    b2.font.color.rgb = RGBColor(30, 64, 175)

    heading = doc.add_heading("Calificacion Final (Agente Juez)", level=1)
    heading.runs[0].font.color.rgb = RGBColor(17, 24, 39)

    for serie in series:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cells = table.rows[0].cells
        cells[0].text = serie["romano"]
        cells[1].text = (
            f"{_clean_text(serie['nombre'])}\n"
            f"{len(serie['items'])} item(s) - {_fmt_score(serie['punteo_obtenido'])} / {_fmt_score(serie['punteo_maximo'])} pts"
        )
        cells[2].text = f"{_fmt_score(serie['punteo_obtenido'])} / {_fmt_score(serie['punteo_maximo'])} pts"
        for cell in cells:
            _shade_cell(cell, "F1F5F9")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.font.size = Pt(9)
        cells[0].paragraphs[0].runs[0].bold = True
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[2].paragraphs[0].runs[0].bold = True

        for index, item in enumerate(serie["items"], start=1):
            correcta = bool(item.get("es_correcta"))
            item_table = doc.add_table(rows=1, cols=2)
            item_table.style = "Table Grid"
            item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            item_left, item_right = item_table.rows[0].cells
            fill = "ECFDF5" if correcta else "FEF2F2"
            for cell in (item_left, item_right):
                _shade_cell(cell, fill)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            status = "Correcto" if correcta else "Incorrecto"
            status_run = item_left.paragraphs[0].add_run(f"Item {index}  {status}")
            status_run.bold = True
            status_run.font.size = Pt(10)
            status_run.font.color.rgb = RGBColor(22, 163, 74) if correcta else RGBColor(220, 38, 38)

            pregunta = _clean_text(item.get("texto_pregunta") or item.get("enunciado"))
            resp = _clean_text(item.get("respuesta_estudiante"))
            just = _clean_text(item.get("justificacion"))

            if pregunta:
                p = item_left.add_paragraph(pregunta)
                p.paragraph_format.space_after = Pt(2)
            if resp:
                p = item_left.add_paragraph()
                p.add_run("Respuesta: ").bold = True
                p.add_run(resp)
                p.paragraph_format.space_after = Pt(2)
            if just:
                p = item_left.add_paragraph(just)
                p.paragraph_format.space_after = Pt(2)

            score_p = item_right.paragraphs[0]
            score_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sr = score_p.add_run(f"{_fmt_score(item.get('punteo_obtenido'))}/{_fmt_score(item.get('punteo_maximo'))}")
            sr.bold = True
            sr.font.size = Pt(14)
            sr.font.color.rgb = RGBColor(22, 163, 74) if correcta else RGBColor(220, 38, 38)
            label = item_right.add_paragraph("puntos")
            label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            label.runs[0].font.size = Pt(8)

        doc.add_paragraph()

    doc.add_heading("Conclusion General", level=1)
    conclusion_table = doc.add_table(rows=1, cols=1)
    conclusion_table.style = "Table Grid"
    conclusion_cell = conclusion_table.rows[0].cells[0]
    _shade_cell(conclusion_cell, "F8FAFC")
    conclusion_cell.text = _clean_text(resultado.get("conclusion") or "Sin conclusion.")

    for title, key in [
        ("Fortalezas", "fortalezas"),
        ("Areas de Mejora", "debilidades"),
        ("Sugerencias", "sugerencias"),
    ]:
        items = resultado.get(key) or []
        if not items:
            continue
        h = doc.add_heading(title, level=2)
        h.runs[0].font.color.rgb = (
            RGBColor(22, 101, 52)
            if key == "fortalezas"
            else RGBColor(185, 28, 28)
            if key == "debilidades"
            else RGBColor(30, 64, 175)
        )
        for item in items:
            doc.add_paragraph(_clean_text(item), style="List Bullet")

    safe_name = re.sub(r"[^\w\-]", "_", estudiante)[:40] or "estudiante"
    suffix = f"_{examen_id}" if examen_id else ""
    filename = f"resultado_{safe_name}{suffix}.docx"
    output_path = str(Path(output_dir) / filename)
    doc.save(output_path)
    return output_path


def generate_word_report(resultado: dict, output_dir: str | None = None) -> str:
    return _generate_word_report_pretty(resultado, output_dir)

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

    output_dir = output_dir or settings.reports_dir
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    estudiante = _clean_text(resultado.get("estudiante") or "Estudiante")
    serie_general = _clean_text(resultado.get("serie") or "Sin serie")
    examen_id = resultado.get("examen_id")
    fecha = resultado.get("fecha") or datetime.now().strftime("%d/%m/%Y, %I:%M %p")
    nivel = resultado.get("nivel_aplicado", 5)
    punteo = _num(resultado.get("punteo_total"))
    maximo = _num(resultado.get("punteo_maximo_total"))
    porcentaje = _num(resultado.get("porcentaje"))
    series = _series_agrupadas(resultado.get("preguntas", []))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9)

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_left, h_right = header.rows[0].cells
    for cell in (h_left, h_right):
        _shade_cell(cell, "0F172A")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    title = h_left.paragraphs[0]
    run = title.add_run(estudiante)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(248, 250, 252)

    meta = h_left.add_paragraph()
    meta_text = serie_general
    if examen_id:
        meta_text += f" · Examen #{examen_id}"
    meta_text += f" · {fecha}"
    meta_run = meta.add_run(meta_text)
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = RGBColor(203, 213, 225)

    score = h_right.paragraphs[0]
    score.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    score_run = score.add_run(f"{porcentaje:.1f}%")
    score_run.bold = True
    score_run.font.size = Pt(22)
    score_run.font.color.rgb = RGBColor(74, 222, 128) if porcentaje >= 70 else RGBColor(248, 113, 113)

    score_pts = h_right.add_paragraph()
    score_pts.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pts_run = score_pts.add_run(f"{_fmt_score(punteo)} / {_fmt_score(maximo)} pts")
    pts_run.font.size = Pt(8)
    pts_run.font.color.rgb = RGBColor(226, 232, 240)

    badges = doc.add_paragraph()
    badges.paragraph_format.space_before = Pt(8)
    b1 = badges.add_run("  completado  ")
    b1.bold = True
    b1.font.color.rgb = RGBColor(22, 101, 52)
    badges.add_run("   ")
    b2 = badges.add_run(f"  Exigencia {nivel}/10  ")
    b2.bold = True
    b2.font.color.rgb = RGBColor(30, 64, 175)

    heading = doc.add_heading("Calificacion Final (Agente Juez)", level=1)
    heading.runs[0].font.color.rgb = RGBColor(17, 24, 39)

    for serie in series:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cells = table.rows[0].cells
        cells[0].text = serie["romano"]
        cells[1].text = (
            f"{_clean_text(serie['nombre'])}\n"
            f"{len(serie['items'])} item(s) · {_fmt_score(serie['punteo_obtenido'])} / {_fmt_score(serie['punteo_maximo'])} pts"
        )
        cells[2].text = f"{_fmt_score(serie['punteo_obtenido'])} / {_fmt_score(serie['punteo_maximo'])} pts"

        for cell in cells:
            _shade_cell(cell, "F1F5F9")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.font.size = Pt(9)
        cells[0].paragraphs[0].runs[0].bold = True
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[2].paragraphs[0].runs[0].bold = True

        for index, item in enumerate(serie["items"], start=1):
            correcta = bool(item.get("es_correcta"))
            item_table = doc.add_table(rows=1, cols=2)
            item_table.style = "Table Grid"
            item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            left, right = item_table.rows[0].cells
            fill = "ECFDF5" if correcta else "FEF2F2"
            for cell in (left, right):
                _shade_cell(cell, fill)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            status = "Correcto" if correcta else "Incorrecto"
            status_run = left.paragraphs[0].add_run(f"Item {index}  {status}")
            status_run.bold = True
            status_run.font.size = Pt(10)
            status_run.font.color.rgb = RGBColor(22, 163, 74) if correcta else RGBColor(220, 38, 38)

            pregunta = _clean_text(item.get("texto_pregunta") or item.get("enunciado"))
            resp = _clean_text(item.get("respuesta_estudiante"))
            just = _clean_text(item.get("justificacion"))

            if pregunta:
                p = left.add_paragraph(pregunta)
                p.paragraph_format.space_after = Pt(2)
            if resp:
                p = left.add_paragraph()
                p.add_run("Respuesta: ").bold = True
                p.add_run(resp)
                p.paragraph_format.space_after = Pt(2)
            if just:
                p = left.add_paragraph(just)
                p.paragraph_format.space_after = Pt(2)

            score_p = right.paragraphs[0]
            score_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sr = score_p.add_run(f"{_fmt_score(item.get('punteo_obtenido'))}/{_fmt_score(item.get('punteo_maximo'))}")
            sr.bold = True
            sr.font.size = Pt(14)
            sr.font.color.rgb = RGBColor(22, 163, 74) if correcta else RGBColor(220, 38, 38)
            label = right.add_paragraph("puntos")
            label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            label.runs[0].font.size = Pt(8)

        doc.add_paragraph()

    doc.add_heading("Conclusion General", level=1)
    conclusion_table = doc.add_table(rows=1, cols=1)
    conclusion_table.style = "Table Grid"
    conclusion_cell = conclusion_table.rows[0].cells[0]
    _shade_cell(conclusion_cell, "F8FAFC")
    conclusion_cell.text = _clean_text(resultado.get("conclusion") or "Sin conclusion.")

    for title, key in [
        ("Fortalezas", "fortalezas"),
        ("Areas de Mejora", "debilidades"),
        ("Sugerencias", "sugerencias"),
    ]:
        items = resultado.get(key) or []
        if not items:
            continue

        h = doc.add_heading(title, level=2)
        h.runs[0].font.color.rgb = (
            RGBColor(22, 101, 52)
            if key == "fortalezas"
            else RGBColor(185, 28, 28)
            if key == "debilidades"
            else RGBColor(30, 64, 175)
        )

        for item in items:
            doc.add_paragraph(_clean_text(item), style="List Bullet")

    safe_name = re.sub(r"[^\w\-]", "_", estudiante)[:40] or "estudiante"
    suffix = f"_{examen_id}" if examen_id else ""
    filename = f"resultado_{safe_name}{suffix}.docx"
    output_path = str(Path(output_dir) / filename)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Version usada por la pantalla de detalle: reporte por series e items.
# Se redefine al final para conservar compatibilidad con imports existentes.
# ---------------------------------------------------------------------------

def _num(value, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _fmt_score(value) -> str:
    n = _num(value)
    return str(int(n)) if n.is_integer() else f"{n:.2f}".rstrip("0").rstrip(".")


def _clean_text(value: object) -> str:
    text = _latex_to_text(str(value or ""))
    text = text.replace("$", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _pdf_safe(value: object) -> str:
    text = _clean_text(value)
    replacements = {
        "—": "-",
        "–": "-",
        "·": "-",
        "•": "-",
        "√": "sqrt",
        "×": "x",
        "÷": "/",
        "≤": "<=",
        "≥": ">=",
        "≠": "!=",
        "≈": "~",
        "∞": "infinito",
        "α": "alpha",
        "β": "beta",
        "π": "pi",
        "✓": "OK",
        "✗": "X",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("latin-1", "replace").decode("latin-1")


def _normalizar_serie(nombre: object) -> str:
    limpio = str(nombre or "").strip()
    if not limpio:
        return "Sin serie"
    if limpio.lower().startswith("serie"):
        return limpio
    if re.fullmatch(r"[IVXLCDM]+", limpio, flags=re.IGNORECASE):
        return f"Serie {limpio.upper()}"
    if limpio.isdigit():
        return f"Serie {limpio}"
    return limpio


def _romano_desde_serie(nombre: str) -> str:
    match = re.search(r"\b(I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII|XIII|XIV|XV)\b", nombre.upper())
    if match:
        return match.group(1)
    match = re.search(r"\d+", nombre)
    return match.group(0) if match else "-"


def _series_agrupadas(preguntas: list[dict]) -> list[dict]:
    grupos: dict[str, dict] = {}

    for pregunta in preguntas or []:
        nombre = _normalizar_serie(
            pregunta.get("serie_seccion")
            or pregunta.get("serie")
            or pregunta.get("serie_pregunta")
        )
        if nombre not in grupos:
            grupos[nombre] = {
                "nombre": nombre,
                "romano": _romano_desde_serie(nombre),
                "items": [],
            }
        grupos[nombre]["items"].append(pregunta)

    series = []
    for serie in grupos.values():
        items = sorted(serie["items"], key=lambda p: _num(p.get("numero"), 9999))
        obtenido = sum(_num(p.get("punteo_obtenido")) for p in items)
        maximo = sum(_num(p.get("punteo_maximo")) for p in items)
        series.append({
            **serie,
            "items": items,
            "punteo_obtenido": obtenido,
            "punteo_maximo": maximo,
        })

    return series


def generate_pdf_report(resultado: dict, output_dir: str | None = None) -> str:
    from fpdf import FPDF

    output_dir = output_dir or settings.reports_dir
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    estudiante = _clean_text(resultado.get("estudiante") or "Estudiante")
    serie_general = _clean_text(resultado.get("serie") or "Sin serie")
    examen_id = resultado.get("examen_id")
    fecha = resultado.get("fecha") or datetime.now().strftime("%d/%m/%Y, %I:%M %p")
    nivel = resultado.get("nivel_aplicado", 5)
    punteo = _num(resultado.get("punteo_total"))
    maximo = _num(resultado.get("punteo_maximo_total"))
    porcentaje = _num(resultado.get("porcentaje"))
    series = _series_agrupadas(resultado.get("preguntas", []))

    class ResultPDF(FPDF):
        def cell(self, w=0, h=0, txt="", border=0, ln=0, align="", fill=False, link=""):
            return super().cell(w, h, _pdf_safe(txt), border, ln, align, fill, link)

        def multi_cell(self, w, h, txt="", border=0, align="J", fill=False, split_only=False):
            return super().multi_cell(w, h, _pdf_safe(txt), border, align, fill, split_only)

        def header(self):
            self.set_fill_color(15, 23, 42)
            self.rect(0, 0, 210, 29, "F")
            self.set_fill_color(34, 197, 94)
            self.rect(0, 0, 5, 29, "F")
            self.set_text_color(255, 255, 255)
            self.set_font("Helvetica", "B", 15)
            self.set_xy(12, 7)
            self.cell(120, 8, "Informe de Calificacion")
            self.set_xy(12, 17)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(203, 213, 225)
            self.cell(100, 5, "Resultado final por series e items")
            self.set_font("Helvetica", "", 8)
            self.set_xy(150, 8)
            self.set_text_color(226, 232, 240)
            self.cell(48, 5, "Generado por DeepGrader AI", align="R")

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(120, 120, 120)
            self.cell(0, 6, f"Pagina {self.page_no()}", align="C")

    pdf = ResultPDF()
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.add_page()
    pdf.set_y(38)

    pdf.set_fill_color(248, 250, 252)
    pdf.set_draw_color(226, 232, 240)
    pdf.rect(10, 36, 190, 35, "DF")

    pdf.set_text_color(17, 24, 39)
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_xy(16, 43)
    pdf.cell(130, 9, estudiante, ln=0)

    score_color = (22, 163, 74) if porcentaje >= 70 else (202, 138, 4) if porcentaje >= 60 else (220, 38, 38)
    pdf.set_fill_color(255, 255, 255)
    pdf.set_draw_color(*score_color)
    pdf.rect(151, 41, 40, 23, "DF")
    pdf.set_xy(151, 45)
    pdf.set_text_color(*score_color)
    pdf.set_font("Helvetica", "B", 17)
    pdf.cell(40, 8, f"{porcentaje:.1f}%", align="C")
    pdf.set_xy(151, 55)
    pdf.set_text_color(75, 85, 99)
    pdf.set_font("Helvetica", "", 8)
    pdf.cell(40, 4, f"{_fmt_score(punteo)} / {_fmt_score(maximo)} pts", align="C")

    pdf.set_xy(16, 55)
    pdf.set_text_color(55, 65, 81)
    pdf.set_font("Helvetica", "", 9)
    meta = f"{serie_general}"
    if examen_id:
        meta += f" · Examen #{examen_id}"
    meta += f" · {fecha}"
    pdf.multi_cell(130, 5, meta)

    pdf.set_y(76)
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_fill_color(236, 253, 245)
    pdf.set_text_color(22, 101, 52)
    pdf.cell(28, 7, " completado ", border=1, ln=0, align="C", fill=True)
    pdf.set_x(44)
    pdf.set_fill_color(239, 246, 255)
    pdf.set_text_color(30, 64, 175)
    pdf.cell(36, 7, f" Exigencia {nivel}/10 ", border=1, ln=True, align="C", fill=True)

    pdf.ln(8)
    pdf.set_text_color(17, 24, 39)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, "Calificacion Final (Agente Juez)", ln=True)
    pdf.set_draw_color(34, 197, 94)
    pdf.line(12, pdf.get_y(), 198, pdf.get_y())
    pdf.ln(3)

    for serie in series:
        if pdf.get_y() > 240:
            pdf.add_page()
            pdf.set_y(34)

        y = pdf.get_y() + 3
        pdf.set_fill_color(241, 245, 249)
        pdf.set_draw_color(203, 213, 225)
        pdf.rect(12, y, 186, 20, "DF")
        pdf.set_xy(18, y + 5)
        pdf.set_fill_color(220, 252, 231)
        pdf.set_text_color(77, 124, 15)
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(12, 9, serie["romano"], border=1, align="C", fill=True)
        pdf.set_xy(35, y + 4)
        pdf.set_text_color(17, 24, 39)
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(100, 5, _clean_text(serie["nombre"]), ln=True)
        pdf.set_x(35)
        pdf.set_text_color(75, 85, 99)
        pdf.set_font("Helvetica", "", 8)
        pdf.cell(
            100,
            5,
            f"{len(serie['items'])} item(s) · {_fmt_score(serie['punteo_obtenido'])} / {_fmt_score(serie['punteo_maximo'])} pts",
        )
        pdf.set_xy(162, y + 6)
        pdf.set_text_color(77, 124, 15)
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(28, 6, f"{_fmt_score(serie['punteo_obtenido'])} / {_fmt_score(serie['punteo_maximo'])} pts", border=1, align="C")
        pdf.set_y(y + 26)

        for index, item in enumerate(serie["items"], start=1):
            correcta = bool(item.get("es_correcta"))
            item_h = 35
            just = _clean_text(item.get("justificacion"))
            resp = _clean_text(item.get("respuesta_estudiante"))
            pregunta = _clean_text(item.get("texto_pregunta") or item.get("enunciado"))
            if just:
                item_h += min(16, max(0, len(just) // 95) * 5)
            if pregunta:
                item_h += 5
            if pdf.get_y() + item_h > 275:
                pdf.add_page()
                pdf.set_y(34)

            y = pdf.get_y()
            if correcta:
                pdf.set_fill_color(240, 253, 244)
                pdf.set_draw_color(187, 247, 208)
                title_color = (22, 163, 74)
            else:
                pdf.set_fill_color(254, 242, 242)
                pdf.set_draw_color(254, 202, 202)
                title_color = (220, 38, 38)
            pdf.rect(18, y, 174, item_h, "DF")
            pdf.set_xy(24, y + 6)
            pdf.set_text_color(*title_color)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(70, 5, f"Item {index} - {'Correcto' if correcta else 'Incorrecto'}")
            pdf.set_xy(164, y + 6)
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(22, 7, f"{_fmt_score(item.get('punteo_obtenido'))}/{_fmt_score(item.get('punteo_maximo'))}", align="R")
            pdf.set_xy(24, y + 15)
            pdf.set_text_color(17, 24, 39)
            pdf.set_font("Helvetica", "", 8)
            if pregunta:
                pdf.multi_cell(130, 4, pregunta[:220])
                pdf.set_x(24)
            if resp:
                pdf.set_font("Helvetica", "B", 8)
                pdf.cell(20, 4, "Respuesta: ")
                pdf.set_font("Helvetica", "", 8)
                pdf.multi_cell(130, 4, resp[:260])
                pdf.set_x(24)
            if just:
                pdf.set_font("Helvetica", "", 8)
                pdf.multi_cell(140, 4, just[:500])
            pdf.set_y(y + item_h + 4)

    conclusion = _clean_text(resultado.get("conclusion") or "Sin conclusion.")
    pdf.ln(2)
    if pdf.get_y() > 220:
        pdf.add_page()
        pdf.set_y(34)
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(17, 24, 39)
    pdf.cell(0, 7, "Conclusion General", ln=True)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_x(12)
    pdf.multi_cell(186, 5, conclusion)

    for title, key in [
        ("Fortalezas", "fortalezas"),
        ("Areas de Mejora", "debilidades"),
        ("Sugerencias", "sugerencias"),
    ]:
        items = resultado.get(key) or []
        if not items:
            continue
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 6, title, ln=True)
        pdf.set_font("Helvetica", "", 9)
        for item in items:
            pdf.set_x(12)
            pdf.multi_cell(186, 5, f"- {_clean_text(item)}")

    safe_name = re.sub(r"[^\w\-]", "_", estudiante)[:40] or "estudiante"
    suffix = f"_{examen_id}" if examen_id else ""
    filename = f"resultado_{safe_name}{suffix}.pdf"
    output_path = str(Path(output_dir) / filename)
    pdf.output(output_path)
    return output_path


def _shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def generate_word_report(resultado: dict, output_dir: str | None = None) -> str:
    return _generate_word_report_pretty(resultado, output_dir)

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    output_dir = output_dir or settings.reports_dir
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    estudiante = _clean_text(resultado.get("estudiante") or "Estudiante")
    serie_general = _clean_text(resultado.get("serie") or "Sin serie")
    examen_id = resultado.get("examen_id")
    fecha = resultado.get("fecha") or datetime.now().strftime("%d/%m/%Y, %I:%M %p")
    nivel = resultado.get("nivel_aplicado", 5)
    punteo = _num(resultado.get("punteo_total"))
    maximo = _num(resultado.get("punteo_maximo_total"))
    porcentaje = _num(resultado.get("porcentaje"))
    series = _series_agrupadas(resultado.get("preguntas", []))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title_run = title.add_run(estudiante)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(17, 24, 39)

    meta = doc.add_paragraph()
    meta_text = f"{serie_general}"
    if examen_id:
        meta_text += f" · Examen #{examen_id}"
    meta_text += f" · {fecha}"
    meta.add_run(meta_text).font.size = Pt(9)

    score = doc.add_paragraph()
    score.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = score.add_run(f"{porcentaje:.1f}%")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(22, 163, 74) if porcentaje >= 70 else RGBColor(220, 38, 38)
    score.add_run(f"\n{_fmt_score(punteo)} / {_fmt_score(maximo)} pts")

    badges = doc.add_paragraph()
    badges.add_run("completado").bold = True
    badges.add_run(f"    Exigencia {nivel}/10")

    heading = doc.add_heading("Calificacion Final (Agente Juez)", level=1)
    heading.runs[0].font.color.rgb = RGBColor(17, 24, 39)

    for serie in series:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        cells = table.rows[0].cells
        cells[0].text = serie["romano"]
        cells[1].text = (
            f"{_clean_text(serie['nombre'])}\n"
            f"{len(serie['items'])} item(s) · {_fmt_score(serie['punteo_obtenido'])} / {_fmt_score(serie['punteo_maximo'])} pts"
        )
        cells[2].text = f"{_fmt_score(serie['punteo_obtenido'])} / {_fmt_score(serie['punteo_maximo'])} pts"
        for cell in cells:
            _shade_cell(cell, "F9FAFB")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        cells[0].paragraphs[0].runs[0].bold = True
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[2].paragraphs[0].runs[0].bold = True

        for index, item in enumerate(serie["items"], start=1):
            correcta = bool(item.get("es_correcta"))
            item_table = doc.add_table(rows=1, cols=2)
            item_table.style = "Table Grid"
            left, right = item_table.rows[0].cells
            fill = "ECFDF5" if correcta else "FEF2F2"
            _shade_cell(left, fill)
            _shade_cell(right, fill)

            p = left.paragraphs[0]
            status = "Correcto" if correcta else "Incorrecto"
            r = p.add_run(f"Item {index}  {status}")
            r.bold = True
            r.font.color.rgb = RGBColor(22, 163, 74) if correcta else RGBColor(220, 38, 38)

            pregunta = _clean_text(item.get("texto_pregunta") or item.get("enunciado"))
            resp = _clean_text(item.get("respuesta_estudiante"))
            just = _clean_text(item.get("justificacion"))

            if pregunta:
                left.add_paragraph(pregunta)
            if resp:
                rp = left.add_paragraph()
                rp.add_run("Respuesta: ").bold = True
                rp.add_run(resp)
            if just:
                left.add_paragraph(just)

            score_p = right.paragraphs[0]
            score_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            score_run = score_p.add_run(f"{_fmt_score(item.get('punteo_obtenido'))}/{_fmt_score(item.get('punteo_maximo'))}")
            score_run.bold = True
            score_run.font.size = Pt(14)
            score_run.font.color.rgb = RGBColor(22, 163, 74) if correcta else RGBColor(220, 38, 38)
            right.add_paragraph("puntos").alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()

    doc.add_heading("Conclusion General", level=1)
    doc.add_paragraph(_clean_text(resultado.get("conclusion") or "Sin conclusion."))

    for title, key in [
        ("Fortalezas", "fortalezas"),
        ("Areas de Mejora", "debilidades"),
        ("Sugerencias", "sugerencias"),
    ]:
        items = resultado.get(key) or []
        if not items:
            continue
        doc.add_heading(title, level=2)
        for item in items:
            doc.add_paragraph(_clean_text(item), style="List Bullet")

    safe_name = re.sub(r"[^\w\-]", "_", estudiante)[:40] or "estudiante"
    suffix = f"_{examen_id}" if examen_id else ""
    filename = f"resultado_{safe_name}{suffix}.docx"
    output_path = str(Path(output_dir) / filename)
    doc.save(output_path)
    return output_path
